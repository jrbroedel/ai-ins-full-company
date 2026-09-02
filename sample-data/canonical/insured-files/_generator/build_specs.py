#!/usr/bin/env python3
"""Assemble per-insured fill specs for the 3 multi-vehicle insureds.
Values come from (a) the shipped filled submissions (insured constants, policy-level
selections, original vehicle) and (b) frozen canonical donor rows per the adoption map.
No value here is rated or invented beyond presentation conventions already used in the
shipped set (garage/security/usage/mileage strings, deductible-by-class rule)."""
import json, datetime as dt, re
import openpyxl
from docx import Document

WS = "/home/claude/ws/01_Completed_Files"
DS = json.load(open("/home/claude/repo/sample-data/canonical/canonical_dataset.json"))
MAP = json.load(open("vehicle_adoption_map.json"))
SUBS = {s["seq"]: s for s in DS["submissions"]}

RATER = "/home/claude/repo/templates/Exotic_Collector_Auto_MGA_Rating_Matrix_7.xlsx"
wb = openpyxl.load_workbook(RATER, read_only=True, data_only=True)
CLASSES = {}
for row in wb["Base Rates"].iter_rows():
    for cell in row:
        v = cell.value
        if isinstance(v, str) and re.match(r"^\d{2} - ", v):
            CLASSES[int(v[:2])] = v
wb.close()
assert len(CLASSES) >= 12, CLASSES

STATEFULL = {"KS": "Kansas", "CA": "California", "MA": "Massachusetts"}
METRO = {"KS": "Kansas City, MO-KS", "CA": "Los Angeles - Long Beach, CA", "MA": "Boston metro, MA"}
SLRATE = {"KS": 0.06, "CA": 0.03, "MA": 0.04}


RATER_MAKE = {"BMW": "BMW / BMW M", "Mercedes-Benz": None}  # None -> model-dependent below
def rater_make(make, model):
    if make == "Mercedes-Benz":
        return "Mercedes-AMG" if "AMG" in model else "Mercedes-Maybach" if "Maybach" in model else "Mercedes-AMG" if False else "Mercedes-AMG" if "GT" in model else "Other / Bespoke"
    return RATER_MAKE.get(make, make)

YEAR_FIX = {"911 Reimagined": 2022, "GranTurismo": 2023, "8C 2900": 1938,
            "Continental GT": 2023, "Huracan EVO": 2021, "750S": 2024,
            "M8 Competition": 2023, "SF90 Stradale": 2023, "AMG GT": 2021,
            "DB5": 1964, "296 GTB": 2023, "4.5 Litre": 1929, "911 Turbo S": 2022,
            "DB12": 2024, "280SL Pagoda": 1967, "E-Type Series 1": 1963,
            "275 GTB/4": 1967}
def fix_year(model, year):
    for k, y in YEAR_FIX.items():
        if k.lower() in model.lower(): return y
    return year

def coll_band(n):
    for lo, hi, lbl in [(1,1,"1 vehicle"),(2,3,"2 - 3 vehicles"),(4,6,"4 - 6 vehicles"),
                        (7,10,"7 - 10 vehicles"),(11,20,"11 - 20 vehicles")]:
        if lo <= n <= hi: return lbl
    return "21 + vehicles"

def shipped_constants(folder):
    d = Document(f"{WS}/{folder}/01_Submission_{folder.split(', ')[0]}_{folder.split(', ')[1]}.docx")
    T = d.tables
    cell = lambda t,c: T[t].rows[0].cells[c].text.strip()
    office_raw = T[1].rows[0].cells[0].text
    ref = re.search(r"Submission Ref: (\S+)", office_raw).group(1)
    dr = re.search(r"Date Received: ([A-Za-z]+ \d+, \d{4})", office_raw).group(1)
    checked = {}
    for p in d.paragraphs:
        if "\u2612" in p.text:
            for m in re.finditer(r"\u2612 ([^\u2610\u2612]+?)(?:      |$)", p.text):
                checked.setdefault(p.text[:20], []).append(m.group(1).strip())
    def sel(header_frag):
        for i, p in enumerate(d.paragraphs):
            if p.text.strip().startswith(header_frag):
                nxt = d.paragraphs[i+1].text
                m = re.search(r"\u2612 ([^\u2610]+?)(?:    |$)", nxt)
                return m.group(1).strip() if m else None
    drv = T[13].rows[1]
    return {
        "insured": {"broker_firm": cell(2,1), "broker_license": cell(2,3),
                    "broker_contact": cell(3,1), "broker_email": cell(3,3),
                    "broker_phone": cell(4,1), "broker_address": cell(4,3),
                    "name": cell(5,1), "address": cell(6,1), "email": cell(7,1),
                    "phone": cell(7,3), "effective": cell(8,1), "expiry": cell(8,3)},
        "office": {"ref": ref, "date_received": dr, "disposition": "Bound"},
        "sel": {"entity": sel("ENTITY TYPE"), "policy_type": sel("POLICY TYPE"),
                "vclass_v1": sel("VEHICLE CLASS"), "valuation": sel("VALUATION BASIS"),
                "storage": sel("STORAGE / GARAGING"), "security": sel("SECURITY / ANTI-THEFT"),
                "usage": sel("USAGE CATEGORY \u2014 PERSONAL"), "mileage": sel("ANNUAL MILEAGE BAND"),
                "age_band": sel("PRINCIPAL DRIVER AGE BAND"), "hpde": sel("HIGH-PERFORMANCE DRIVING EXPERIENCE"),
                "mvr": sel("MOTOR VEHICLE RECORD"), "claims": sel("CLAIMS RECORD"),
                "csl": sel("LIABILITY LIMIT"), "pd_deductible": sel("PHYSICAL DAMAGE DEDUCTIBLE"),
                "deductible_basis": sel("OPTIONAL DEDUCTIBLE BASES"), "attachments": sel("INCLUDED WITH THIS SUBMISSION")},
        "vehicle1": {"year": cell(9,1), "make": cell(9,3), "model": cell(9,5),
                     "vin": cell(10,1), "agreed_value_disp": cell(10,3),
                     "garaging_address": cell(11,1), "garaging_state": cell(12,1),
                     "garaging_zip": cell(12,3)},
        "driver": {"name": drv.cells[0].text.strip(), "dob": drv.cells[1].text.strip(),
                   "license": drv.cells[2].text.strip(), "hp_years": drv.cells[3].text.strip()},
    }

def donor_vehicle(seq, base):
    s = SUBS[seq]; v = s["vehicle"]; cls_no = s["quote"]["rating_basis"]["rating_vehicle_class"]
    cls = CLASSES[cls_no]
    make = "Singer Vehicle Design" if "singer" in v["make"].lower() else v["make"]
    ded = "$5,000" if cls_no in (3,4,5) else "$2,500"
    year = fix_year(v["model"], v["year"])
    return {"seq": seq, "year": str(year), "make": v["make"], "model": v["model"],
            "rater_make": rater_make(make, v["model"]),
            "vin": v["vin"], "agreed_value": v["agreed_value"],
            "agreed_value_disp": f"${v['agreed_value']:,}",
            "garaging_address": base["vehicle1"]["garaging_address"],
            "garaging_state": base["vehicle1"]["garaging_state"],
            "garaging_zip": base["vehicle1"]["garaging_zip"],
            "vclass": cls, "vclass_no": cls_no, "valuation": base["sel"]["valuation"],
            "storage": base["sel"]["storage"], "security": base["sel"]["security"],
            "usage": base["sel"]["usage"], "mileage": base["sel"]["mileage"],
            "pd_deductible_row": ded,
            "effective_date": s["policy"]["effective_date"],
            "bound_premium": s["policy"]["money"]["bound_premium"],
            "policy_number": s["policy"]["policy_number"],
            "policy_fee": 250.0, "inspection_fee": 150.0 if v["agreed_value"] >= 250000 else 0.0}

def main():
    specs = {}
    for key, m in MAP.items():
        last, first = key.split(", ")
        folder = key
        base = shipped_constants(folder)
        orig = SUBS[m["original_seq"]]
        v1 = dict(donor_vehicle(m["original_seq"], base))
        # vehicle 1 keeps its shipped display values (year etc. were session-fixed)
        v1.update(base["vehicle1"])
        v1["agreed_value"] = orig["vehicle"]["agreed_value"]
        v1["vclass"] = base["sel"]["vclass_v1"]
        v1["vclass_no"] = int(base["sel"]["vclass_v1"][:2])
        v1["pd_deductible_row"] = "$5,000" if v1["vclass_no"] in (3,4,5) else "$2,500"
        vehicles = [v1] + [donor_vehicle(a["seq"], base) for a in m["added"]]
        st = m["state"]
        spec = {
            "insured": base["insured"], "office": dict(base["office"]),
            "driver": base["driver"], "state": st, "state_full": STATEFULL[st],
            "metro": METRO[st], "sl_rate": SLRATE[st],
            "policy_selections": {
                "entity": base["sel"]["entity"], "policy_type": base["sel"]["policy_type"],
                "collection_size": coll_band(len(vehicles)),
                "age_band": base["sel"]["age_band"], "hpde": base["sel"]["hpde"],
                "mvr": base["sel"]["mvr"], "claims": base["sel"]["claims"],
                "csl": base["sel"]["csl"], "pd_deductible": base["sel"]["pd_deductible"],
                "deductible_basis": base["sel"]["deductible_basis"],
                "attachments": base["sel"]["attachments"]},
            "uw_answers": {}, "vehicles": vehicles,
        }
        total = sum(v["bound_premium"] for v in vehicles)
        spec["office"]["base_premium"] = f"{total:,.2f}"
        spec["office"]["policy_no"] = orig["policy"]["policy_number"]
        specs[key] = spec
        json.dump(spec, open(f"spec_{last}.json", "w"), indent=1)
        print(f"{key}: {len(vehicles)} vehicles, schedule GWP {total:,.2f}")
    return specs

if __name__ == "__main__":
    main()
