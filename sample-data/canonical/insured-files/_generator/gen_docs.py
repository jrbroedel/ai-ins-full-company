#!/usr/bin/env python3
"""Quote / Binder / Invoice for the three multi-vehicle insureds. python-docx -> LibreOffice
PDF (the shipped set's own toolchain). All money is frozen canonical bound premium; the
coverage split, fee schedule and per-state surplus-lines tax bases replicate the shipped
five exactly (ratios and bases derived from, and verified against, those files)."""
import json, os, subprocess, sys, datetime as dt
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

RED = RGBColor(0xC0, 0x00, 0x00)
GREY = RGBColor(0x59, 0x59, 0x59)
R_L, R_P, R_U = 0.211679, 0.708029, 0.056205
SYNDICATES = [1101, 1408, 1776, 2010, 2288, 2517, 3003, 3621, 4141, 4890]
FOOT = ("Torque Underwriters LLC  |  801 Brickell Key Drive, Suite 1400, Miami, FL 33131"
        "  |  (305) 555-0180  |  bind@torqueuw.example")
SPECIMEN = ("Specimen document \u2014 synthetic demonstration data (ADR 0045 canonical book). "
            "Not a real applicant, policy, or bound risk.")
OUT = "out/docs"; os.makedirs(OUT, exist_ok=True)

def fmtd(d): return dt.date.fromisoformat(d).strftime("%B %-d, %Y") if isinstance(d, str) else d.strftime("%B %-d, %Y")
def money(x): return f"${x:,.2f}"

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    el = tcPr.makeelement(qn("w:shd"), {qn("w:val"): "clear", qn("w:fill"): hexcolor})
    tcPr.append(el)

def para(d, text, size=9, bold=False, color=None, align=None, space_after=4, italic=False):
    p = d.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size); r.font.bold = bold; r.font.italic = italic
    if color: r.font.color.rgb = color
    if align: p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p

def masthead(d, doctype):
    t = d.add_table(rows=1, cols=2); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.columns[0].width = Inches(4.6); t.columns[1].width = Inches(2.4)
    c = t.rows[0].cells[0].paragraphs[0]
    r1 = c.add_run("TORQUE "); r1.font.size = Pt(16); r1.font.bold = True
    r2 = c.add_run("UNDERWRITERS"); r2.font.size = Pt(16); r2.font.bold = True; r2.font.color.rgb = RED
    sub = t.rows[0].cells[0].add_paragraph()
    rs = sub.add_run("Coverholder at Lloyd's   |   Exotic & Collector Automobile Program")
    rs.font.size = Pt(8); rs.font.color.rgb = GREY
    c2 = t.rows[0].cells[1].paragraphs[0]; c2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r3 = c2.add_run(doctype); r3.font.size = Pt(14); r3.font.bold = True
    para(d, "", size=2, space_after=2)

def refgrid(d, pairs):
    t = d.add_table(rows=2, cols=len(pairs))
    for i, (lbl, val) in enumerate(pairs):
        p = t.rows[0].cells[i].paragraphs[0]
        r = p.add_run(lbl); r.font.size = Pt(6.5); r.font.bold = True; r.font.color.rgb = GREY
        p2 = t.rows[1].cells[i].paragraphs[0]
        r2 = p2.add_run(val); r2.font.size = Pt(9); r2.font.bold = True
    para(d, "", size=2, space_after=4)

def kv(d, rows, lw=1.7):
    t = d.add_table(rows=len(rows), cols=2)
    t.columns[0].width = Inches(lw); t.columns[1].width = Inches(7.0 - lw)
    for i, (lbl, val) in enumerate(rows):
        p = t.rows[i].cells[0].paragraphs[0]
        r = p.add_run(lbl); r.font.size = Pt(8); r.font.bold = True; r.font.color.rgb = GREY
        p2 = t.rows[i].cells[1].paragraphs[0]
        r2 = p2.add_run(val); r2.font.size = Pt(9)
    para(d, "", size=2, space_after=4)

def head(d, text):
    para(d, text, size=11, bold=True, space_after=3)

def table_grid(t):
    tbl = t._tbl
    pr = tbl.tblPr
    el = pr.makeelement(qn("w:tblBorders"), {})
    for edge in ("top","left","bottom","right","insideH","insideV"):
        e = el.makeelement(qn(f"w:{edge}"), {qn("w:val"): "single", qn("w:sz"): "4",
                                             qn("w:color"): "BBBBBB"})
        el.append(e)
    pr.append(el)

def schedule_table(d, spec, certs, name, with_premium=True):
    cols = ["#", "Year / Make / Model", "VIN", "Certificate",
            "Agreed Value"] + (["Annual Premium"] if with_premium else [])
    t = d.add_table(rows=1 + len(spec["vehicles"]), cols=len(cols)); table_grid(t)
    for i, cname in enumerate(cols):
        p = t.rows[0].cells[i].paragraphs[0]
        r = p.add_run(cname); r.font.size = Pt(7.5); r.font.bold = True
        shade(t.rows[0].cells[i], "F2F2F2")
    for vi, v in enumerate(spec["vehicles"]):
        cells = t.rows[vi + 1].cells
        vals = [str(vi + 1), f"{v['year']} {v['make']} {v['model']}", v["vin"],
                f"TQ-C-2026-{certs[f'{name},{vi}']:04d}",
                f"${v['agreed_value']:,}"]
        if with_premium: vals.append(money(v["bound_premium"]))
        for ci, val in enumerate(vals):
            p = cells[ci].paragraphs[0]; r = p.add_run(val); r.font.size = Pt(7.5)
    para(d, "", size=2, space_after=4)

def money_rows(d, rows, total_label, total_val):
    t = d.add_table(rows=len(rows) + 1, cols=2); table_grid(t)
    t.columns[0].width = Inches(5.4); t.columns[1].width = Inches(1.6)
    for i, (lbl, val) in enumerate(rows):
        p = t.rows[i].cells[0].paragraphs[0]; r = p.add_run(lbl); r.font.size = Pt(9)
        p2 = t.rows[i].cells[1].paragraphs[0]; p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r2 = p2.add_run(money(val)); r2.font.size = Pt(9)
    p = t.rows[-1].cells[0].paragraphs[0]; r = p.add_run(total_label); r.font.size = Pt(9); r.font.bold = True
    p2 = t.rows[-1].cells[1].paragraphs[0]; p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r2 = p2.add_run(money(total_val)); r2.font.size = Pt(9); r2.font.bold = True
    shade(t.rows[-1].cells[0], "F2F2F2"); shade(t.rows[-1].cells[1], "F2F2F2")
    para(d, "", size=2, space_after=4)

def bullets(d, items):
    for it in items:
        para(d, "\u2022  " + it, size=8, space_after=2)

def footer_block(d):
    para(d, SPECIMEN, size=7, italic=True, color=GREY, space_after=1)
    para(d, FOOT, size=7, color=GREY, space_after=0)

def taxes(spec, prem, fees):
    st = spec["state"]; sf = spec["state_full"]; rate = spec["sl_rate"]
    base = prem + fees
    rows = [(f"{sf} Surplus Lines Tax ({rate*100:.2f}%)", round(base * rate, 2))]
    if st == "CA":
        rows.append(("California SLA Stamping Fee (0.18%)", round(base * 0.0018, 2)))
    return rows

def split(total):
    l = round(total * R_L, 2); p = round(total * R_P, 2); u = round(total * R_U, 2)
    return l, p, u, round(total - l - p - u, 2)

def build_all(name):
    spec = json.load(open(f"spec_{name}.json"))
    certs = json.load(open("cert_map.json"))
    n = len(spec["vehicles"])
    total = round(sum(v["bound_premium"] for v in spec["vehicles"]), 2)
    fees_policy = 250.0 * n
    insp_n = sum(1 for v in spec["vehicles"] if v["agreed_value"] >= 250000)
    fees_insp = 150.0 * insp_n
    fees = fees_policy + fees_insp
    tax_rows = taxes(spec, total, fees)
    grand = round(total + fees + sum(v for _, v in tax_rows), 2)
    ins = spec["insured"]; st_full = spec["state_full"]
    num = {"Delacroix": 147, "Vasquez": 148, "Harrington": 149}[name]
    qref, cref, iref = (f"TQ-{k}-2026-{num:04d}" for k in ("Q", "C", "INV"))
    umr = f"B0999TQ26{num:04d}"
    eff = dt.datetime.strptime(ins["effective"], "%B %d, %Y").date()
    qdate = dt.datetime.strptime(spec["office"]["date_received"], "%B %d, %Y").date()
    issue = eff - dt.timedelta(days=3)
    first = ins["name"].split()[0]
    sl_notice = (f"Surplus Lines Notice ({st_full}): This insurance is placed pursuant to the "
                 f"{st_full} surplus lines law with an insurer not licensed by {st_full}. Surplus "
                 f"lines policies are not protected by the {st_full} insurance guaranty association, "
                 f"and the rates and forms are not approved by any {st_full} regulatory authority.")
    l, p_, u, m = split(total)

    def newdoc():
        d = Document()
        st = d.styles["Normal"].font; st.name = "Calibri"; st.size = Pt(9)
        for s in d.sections:
            s.left_margin = s.right_margin = Inches(0.7)
            s.top_margin = Inches(0.6); s.bottom_margin = Inches(0.6)
        return d

    # ------------- QUOTE -------------
    d = newdoc(); masthead(d, "Quotation")
    refgrid(d, [("QUOTE REFERENCE", qref), ("QUOTE DATE", fmtd(qdate)),
                ("VALID UNTIL", fmtd(qdate + dt.timedelta(days=30))),
                ("UNDERWRITER", "K. Alvarez, Senior UW")])
    kv(d, [("PROPOSED INSURED", f"{ins['name']}\n{ins['address']}"),
           ("PRODUCING BROKER", f"{ins['broker_firm']} ({spec['state']} Surplus Lines Lic. {ins['broker_license']})"),
           ("PROPOSED PERIOD", f"{ins['effective']} 12:01 a.m. to {ins['expiry']} 12:01 a.m."),
           ("SECURITY", "Certain Underwriters at Lloyd's, London, per Binding Authority Agreement "
                        f"BA-EXO-2026-01 / UMR {umr}, subscribed 100%. Coverage bound hereunder is "
                        "surplus lines insurance.")])
    head(d, f"Scheduled Vehicles ({n})")
    schedule_table(d, spec, certs, name)
    head(d, "Coverage & Premium")
    money_rows(d, [
        ("Automobile Liability \u2014 $1,000,000 Combined Single Limit", l),
        (f"Agreed Value Physical Damage \u2014 ${sum(v['agreed_value'] for v in spec['vehicles']):,} "
         "scheduled agreed value, deductibles per vehicle schedule", p_),
        ("Uninsured / Underinsured Motorist \u2014 $1,000,000 CSL (non-stacked)", u),
        ("Medical Payments \u2014 $10,000 per person", m)],
        "Annual Premium (schedule total)", total)
    head(d, "Premium, Fees & State Charges")
    fee_rows = [("Annual Premium", total), (f"Policy Fees ({n} certificates, fully earned)", fees_policy)]
    if fees_insp:
        fee_rows.append((f"Pre-Bind Inspection Fees ({insp_n} vehicles, agreed value \u2265 $250,000)", fees_insp))
    money_rows(d, fee_rows + tax_rows, "Total Cost of Placement", grand)
    head(d, "Key Terms & Conditions")
    bullets(d, [
        "Physical damage settled on an agreed value basis per the Agreed Value Endorsement; no depreciation on a covered total loss.",
        "OEM parts at a repair facility of the insured's choice, including manufacturer-certified facilities.",
        "Annual mileage limitation of 2,500 miles per scheduled vehicle; odometer declarations required at inception and renewal.",
        "Track, timed-event, and competitive-use exclusion applies. Escorted touring events considered by referral.",
        f"Named driver basis: {ins['name']} and spouse only. Drivers under 30 excluded.",
        "Enclosed-garage warranty at the scheduled garaging address for each vehicle when not in use.",
        "Minimum earned premium: 25%. Policy fees and inspection fees fully earned at inception."])
    head(d, "Subjectivities (to be satisfied prior to binding)")
    bullets(d, [
        "Signed and dated application, including 5-year loss history declaration.",
        "Completed pre-bind photo inspection per scheduled vehicle (minimum 12 photos incl. odometer and VIN plate).",
        "Copy of current registration and driver license for each named driver.",
        "MVRs for all named drivers acceptable to underwriters (no more than 1 minor violation, no majors, 5 years).",
        f"Diligent-effort declination documentation as required for {st_full} surplus lines placement."])
    para(d, "This quotation is an offer of terms only and does not bind coverage. Coverage may be "
            "bound only by written confirmation from Torque Underwriters within its binding "
            "authority. " + sl_notice, size=7.5, space_after=4)
    footer_block(d)
    d.save(f"{OUT}/04_Quote_{name}_{first}.docx")

    # ------------- BINDER -------------
    d = newdoc(); masthead(d, "Binder / Cover Note")
    refgrid(d, [("CERTIFICATE NO.", cref + " (master)"), ("BINDING AUTHORITY REF.", "BA-EXO-2026-01"),
                ("UMR", umr), ("ISSUE DATE", fmtd(issue))])
    para(d, "Evidence of Cover. In consideration of the premium shown below and subject to the terms, "
            "conditions, limitations and exclusions of the policy wording referenced herein, Torque "
            f"Underwriters LLC, as coverholder under Binding Authority Agreement {umr}, hereby confirms "
            "that insurance has been bound with the security shown below, effective as stated. This "
            "cover note is temporary evidence of coverage pending issuance of the policy.", size=8)
    kv(d, [("NAMED INSURED", f"{ins['name']}\n{ins['address']}"),
           ("PRODUCING BROKER", f"{ins['broker_firm']} ({spec['state']} Surplus Lines Lic. {ins['broker_license']})"),
           ("PERIOD OF INSURANCE", f"From {ins['effective']} 12:01 a.m. to {ins['expiry']} 12:01 a.m."),
           ("TYPE OF INSURANCE", "Exotic & Collector Automobile \u2014 Agreed Value Physical Damage, Automobile "
                                 "Liability, UM/UIM and Medical Payments, per program wording TQ-EXO-001 (07/26) "
                                 "and endorsements TQ-AV-01 (Agreed Value), TQ-ML-02 (Mileage Limitation), "
                                 "TQ-TR-03 (Track Use Exclusion)."),
           ("LIMITS / DEDUCTIBLES", "Liability $1,000,000 CSL; Physical Damage at scheduled agreed values, "
                                    "deductibles per vehicle schedule ($2,500; $5,000 classes 03\u201305); "
                                    "UM/UIM $1,000,000 CSL (non-stacked); Medical Payments $10,000."),
           ("PREMIUM", f"Annual schedule premium {money(total)} plus fees and state charges; total cost of "
                       f"placement {money(grand)} as detailed in invoice {iref}."),
           ("CONDITIONS", "Named driver basis; enclosed-garage warranty; 2,500-mile annual limitation per "
                          "vehicle; 25% minimum earned premium; premium payable per invoice terms \u2014 "
                          "non-payment may result in cancellation per policy and applicable law.")])
    head(d, f"Scheduled Vehicles ({n})")
    schedule_table(d, spec, certs, name)
    head(d, "Security Schedule \u2014 Certain Underwriters at Lloyd's, London")
    t = d.add_table(rows=len(SYNDICATES) + 2, cols=2); table_grid(t)
    hr = t.rows[0].cells
    for i, h in enumerate(["Security", "Written Line"]):
        r = hr[i].paragraphs[0].add_run(h); r.font.size = Pt(8); r.font.bold = True
        shade(hr[i], "F2F2F2")
    for i, syn in enumerate(SYNDICATES):
        cs = t.rows[i + 1].cells
        cs[0].paragraphs[0].add_run(f"Lloyd's Syndicate {syn}").font.size = Pt(8)
        cs[1].paragraphs[0].add_run("10.00%").font.size = Pt(8)
    cs = t.rows[-1].cells
    r = cs[0].paragraphs[0].add_run("Total Order Hereon"); r.font.size = Pt(8); r.font.bold = True
    r = cs[1].paragraphs[0].add_run("100.00%"); r.font.size = Pt(8); r.font.bold = True
    para(d, "", size=2, space_after=4)
    head(d, "Notices")
    bullets(d, [sl_notice,
        "Service of Suit: Underwriters will, at the request of the insured, submit to the jurisdiction "
        "of a court of competent jurisdiction within the United States, per the Service of Suit Clause "
        "(USA) in the policy.",
        "Cancellation: As per policy wording and applicable surplus lines requirements; policy fees and "
        "inspection fees are fully earned at inception.",
        "Claims Notification: Report claims to Torque Underwriters Claims, (305) 555-0181, "
        f"claims@torqueuw.example, quoting master certificate number {cref} and the per-vehicle "
        "certificate from the schedule above."])
    para(d, "Signed for and on behalf of the subscribing Underwriters,", size=8, space_after=10)
    para(d, "__________________________________          __________________________________", size=8, space_after=1)
    para(d, "Authorized Signatory, Torque Underwriters LLC                              Date", size=8, space_after=6)
    footer_block(d)
    d.save(f"{OUT}/05_Binder_{name}_{first}.docx")

    # ------------- INVOICE -------------
    d = newdoc(); masthead(d, "Premium Invoice")
    refgrid(d, [("INVOICE NO.", iref), ("INVOICE DATE", fmtd(issue)),
                ("CERTIFICATE NO.", cref + " (master)"), ("DUE DATE", ins["effective"])])
    kv(d, [("BILL TO", f"{ins['name']}\nc/o {ins['broker_firm']} ({spec['state']} Surplus Lines Lic. "
                       f"{ins['broker_license']})\n{ins['broker_address']}"),
           ("RE", f"Schedule of {n} vehicles \u2014 Period {ins['effective']} to {ins['expiry']} \u2014 "
                  f"Binding Authority BA-EXO-2026-01 / UMR {umr}")])
    head(d, "Statement of Amounts Due")
    rows = [(f"Vehicle {vi+1} \u2014 {v['year']} {v['make']} {v['model']} "
             f"(TQ-C-2026-{certs[f'{name},{vi}']:04d})",
             v["bound_premium"]) for vi, v in enumerate(spec["vehicles"])]
    rows.append(("Annual Premium (schedule total)", total))
    rows.append((f"Policy Fees ({n} certificates, fully earned)", fees_policy))
    if fees_insp:
        rows.append((f"Pre-Bind Inspection Fees ({insp_n} vehicles, agreed value \u2265 $250,000)", fees_insp))
    rows += tax_rows
    money_rows(d, rows, "Total Due", grand)
    head(d, "Payment Terms")
    bullets(d, [
        "Payment is due on or before the policy effective date. Coverage is subject to the premium payment condition of the cover "
        "note; non-payment may result in cancellation per policy and applicable law.",
        "Remit by ACH/wire to: Torque Underwriters LLC Premium Trust Account, First Meridian Bank, ABA "
        f"067000000, Account 000123456789, Reference: {iref}. Checks payable to \"Torque Underwriters "
        "LLC Premium Trust\".",
        "Funds are held in a fiduciary premium trust account for the benefit of Underwriters. "
        f"{st_full} surplus lines tax is collected and remitted in accordance with {st_full} law.",
        "Policy fees and pre-bind inspection fees are fully earned at inception and non-refundable. "
        "Minimum earned premium is 25% of annual premium."])
    para(d, "Questions regarding this invoice: accounts@torqueuw.example or (305) 555-0182. Please "
            "quote the invoice number on all correspondence.", size=7.5, space_after=4)
    footer_block(d)
    d.save(f"{OUT}/06_Invoice_{name}_{first}.docx")
    return dict(total=total, fees=fees, grand=grand, files=3)

if __name__ == "__main__":
    for name in (sys.argv[1:] or ["Delacroix", "Vasquez", "Harrington"]):
        r = build_all(name)
        print(name, "premium", money(r["total"]), "grand", money(r["grand"]))
    subprocess.run(["python3", "/mnt/skills/public/pptx/scripts/office/soffice.py", "--headless",
                    "--convert-to", "pdf", *[f"{OUT}/{f}" for f in sorted(os.listdir(OUT)) if f.endswith(".docx")],
                    "--outdir", OUT], capture_output=True)
    print("pdfs done")
