#!/usr/bin/env python3
"""Fill Torque's UW submission form (production template) from a per-insured spec.

Multi-vehicle: clones the Section 4 vehicle block (scheduled vehicle + garaging/security
+ usage) once per vehicle, renumbers headings '4.N Scheduled Vehicle - Vehicle N of M',
renames the in-block Garaging/Usage headings, renumbers downstream sections 5-9, and
rewrites the 'one form per vehicle' instruction to policy-level schedule wording.
Checkboxes are the template's glyph runs; filling flips the target option's box to a
checked glyph in place. All data synthetic (ADR 0045); premiums trace to the frozen
canonical book - this script only inserts values, it never rates.

Usage: fill_submission_form.py <template.docx> <spec.json> <out.docx>
Spec: see build_specs.py - insured block, office block, policy-level selections,
and vehicles[] each carrying text fields + per-vehicle checkbox selections.
"""
import copy, json, sys
from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

CHECKED, UNCHECKED = "\u2612", "\u2610"

def ptext(el):
    return "".join(t.text or "" for t in el.iter(qn("w:t")))

def set_para_text(p_el, new):
    ts = list(p_el.iter(qn("w:t")))
    if not ts:
        return
    ts[0].set(qn("xml:space"), "preserve")
    ts[0].text = new
    for t in ts[1:]:
        t.text = ""

def check_option(p_el, option_label):
    """Flip the box immediately preceding option_label to checked."""
    txt = ptext(p_el)
    needle = UNCHECKED + " " + option_label
    if needle not in txt:
        raise KeyError(f"option not found: {option_label!r} in {txt[:90]!r}")
    set_para_text(p_el, txt.replace(needle, CHECKED + " " + option_label, 1))

def fill(template, spec, out):
    d = Document(template)
    body = d.element.body
    vehicles = spec["vehicles"]
    n = len(vehicles)

    # ---- locate + clone the vehicle block (Section 4 heading .. before COLLECTION SIZE) ----
    kids = list(body)
    start = end = None
    for i, el in enumerate(kids):
        if el.tag == qn("w:p"):
            t = ptext(el).strip()
            if start is None and t.startswith("4.") and "Scheduled Vehicle" in t:
                start = i
            if t.startswith("COLLECTION SIZE"):
                end = i
                break
    block = kids[start:end]
    anchor = kids[end]
    for _ in range(2, n + 1):
        for el in block:
            anchor.addprevious(copy.deepcopy(el))

    # ---- heading renumber + instruction rewrite ----
    veh = 0
    for el in body.iter(qn("w:p")):
        t = ptext(el).strip()
        if t.startswith("4.") and "Scheduled Vehicle" in t:
            veh += 1
            set_para_text(el, f"4.{veh}  Scheduled Vehicle \u2014 Vehicle {veh} of {n}"
                          if n > 1 else "4.  Scheduled Vehicle")
        elif t.startswith("5.") and "Garaging" in t:
            set_para_text(el, f"Garaging & Security \u2014 Vehicle {veh}" if n > 1
                          else "5.  Garaging & Security")
        elif t.startswith("6.") and t.endswith("Usage"):
            set_para_text(el, f"Usage \u2014 Vehicle {veh}" if n > 1 else "6.  Usage")
        elif n > 1 and t.startswith("7."):
            set_para_text(el, "5.  Principal Driver(s)")
        elif n > 1 and t.startswith("8."):
            set_para_text(el, "6.  Loss History \u2014 Prior 5 Years")
        elif n > 1 and t.startswith("9."):
            set_para_text(el, "7.  Coverage Selections")
        elif n > 1 and t.startswith("10."):
            set_para_text(el, "8.  Required Attachments")
        elif n > 1 and t.startswith("11."):
            set_para_text(el, "9.  Declarations & Signature")
        elif t.startswith("Complete one form per"):
            set_para_text(el, "Complete one form per policy; schedule each vehicle in "
                          "Section 4. Fields correspond one-to-one to the Portfolio Register "
                          "risk-input columns (one register row per scheduled vehicle) and the "
                          "submission log, so a completed form can be keyed directly to the "
                          "rating engine without transformation. Selections shown are the "
                          "rateable options; anything outside them is a referral.")

    # ---- office-use block ----
    office = spec["office"]
    cell = d.tables[1].rows[0].cells[0]
    for p in cell.paragraphs:
        t = p.text
        if "Submission Ref" in t:
            set_para_text(p._p, f"Submission Ref: {office['ref']}     "
                          f"Date Received: {office['date_received']}")
        elif "Disposition" in t:
            set_para_text(p._p, "Disposition:  "
                          + (CHECKED if office["disposition"] == "Bound" else UNCHECKED) + " Bound   "
                          + (CHECKED if office["disposition"] == "Declined" else UNCHECKED) + " Declined   "
                          + (CHECKED if office["disposition"] == "Referred" else UNCHECKED) + " Referred")
        elif "Indicative" in t:
            set_para_text(p._p, f"Indicative / Base Premium: ${office['base_premium']}     "
                          f"Policy No. (if bound): {office['policy_no']}     "
                          f"Decline / Referral Reason: {office.get('reason','')}")

    # ---- text tables: label -> value fill (nth occurrence for per-vehicle tables) ----
    def fill_labeled(label, value, occurrence=0):
        seen = 0
        for t in d.tables:
            for row in t.rows:
                cells = row.cells
                for ci, c in enumerate(cells):
                    if c.text.strip() == label and ci + 1 < len(cells):
                        if seen == occurrence:
                            set_para_text(cells[ci + 1].paragraphs[0]._p, str(value))
                            return
                        seen += 1
        raise KeyError(f"label not found (occ {occurrence}): {label}")

    ins = spec["insured"]
    for lbl, key in [("BROKERAGE FIRM", "broker_firm"), ("SURPLUS LINES LICENSE NO.", "broker_license"),
                     ("CONTACT NAME", "broker_contact"), ("EMAIL", "broker_email"),
                     ("PHONE", "broker_phone"), ("MAILING ADDRESS", "broker_address")]:
        fill_labeled(lbl, ins[key], 0)
    fill_labeled("NAMED INSURED (AS TO APPEAR ON POLICY)", ins["name"])
    fill_labeled("MAILING ADDRESS", ins["address"], 1)
    fill_labeled("EMAIL", ins["email"], 1)
    fill_labeled("PHONE", ins["phone"], 1)
    fill_labeled("REQUESTED EFFECTIVE DATE", ins["effective"])
    fill_labeled("REQUESTED EXPIRY DATE", ins["expiry"])

    for vi, v in enumerate(vehicles):
        fill_labeled("YEAR", v["year"], vi)
        fill_labeled("MAKE", v["make"], vi)
        fill_labeled("MODEL", v["model"], vi)
        fill_labeled("VIN / CHASSIS NO.", v["vin"], vi)
        fill_labeled("AGREED VALUE REQUESTED (USD)", v["agreed_value_disp"], vi)
        fill_labeled("GARAGING ADDRESS (STREET, CITY)", v["garaging_address"], vi)
        fill_labeled("GARAGING STATE", v["garaging_state"], vi)
        fill_labeled("GARAGING ZIP", v["garaging_zip"], vi)

    # ---- checkbox groups ----
    def nth_group(header_text, occurrence=0):
        """Return the option paragraph element(s) following the nth header paragraph."""
        seen = 0
        kids2 = list(body.iter(qn("w:p")))
        for i, el in enumerate(kids2):
            if ptext(el).strip() == header_text:
                if seen == occurrence:
                    return kids2[i + 1]
                seen += 1
        raise KeyError(f"group header not found (occ {occurrence}): {header_text}")

    pol = spec["policy_selections"]
    for header, key in [("ENTITY TYPE", "entity"), ("POLICY TYPE", "policy_type"),
                        ("COLLECTION SIZE (VEHICLES ON POLICY)", "collection_size"),
                        ("PRINCIPAL DRIVER AGE BAND", "age_band"),
                        ("HIGH-PERFORMANCE DRIVING EXPERIENCE", "hpde"),
                        ("MOTOR VEHICLE RECORD (3 YEARS, WORST PERMITTED DRIVER)", "mvr"),
                        ("CLAIMS RECORD (CHARGEABLE PERIOD 5 YEARS)", "claims"),
                        ("LIABILITY LIMIT (COMBINED SINGLE LIMIT)", "csl"),
                        ("PHYSICAL DAMAGE DEDUCTIBLE", "pd_deductible"),
                        ("OPTIONAL DEDUCTIBLE BASES", "deductible_basis"),
                        ("INCLUDED WITH THIS SUBMISSION", "attachments")]:
        sel = pol[key]
        for opt in (sel if isinstance(sel, list) else [sel]):
            check_option(nth_group(header), opt)

    for vi, v in enumerate(vehicles):
        for header, key in [("VEHICLE CLASS", "vclass"), ("VALUATION BASIS", "valuation"),
                            ("STORAGE / GARAGING", "storage"), ("SECURITY / ANTI-THEFT", "security"),
                            ("USAGE CATEGORY \u2014 PERSONAL", "usage"),
                            ("ANNUAL MILEAGE BAND", "mileage")]:
            check_option(nth_group(header, vi), v[key])

    # ---- driver table + underwriting Y/N ----
    for t in d.tables:
        if t.rows and t.rows[0].cells[0].text.strip() == "Driver name":
            r = t.rows[1]
            drv = spec["driver"]
            for ci, val in enumerate([drv["name"], drv["dob"], drv["license"], drv["hp_years"]]):
                set_para_text(r.cells[ci].paragraphs[0]._p, val)
        if t.rows and t.rows[0].cells[0].text.strip() == "Underwriting question":
            for row in t.rows[1:]:
                c = row.cells[1]
                ans = spec["uw_answers"].get(row.cells[0].text.strip(), "N")
                txt = c.text.replace(UNCHECKED + " " + ans, CHECKED + " " + ans, 1)
                set_para_text(c.paragraphs[0]._p, txt)

    d.save(out)
    return out

if __name__ == "__main__":
    tpl, specf, out = sys.argv[1], sys.argv[2], sys.argv[3]
    fill(tpl, json.load(open(specf)), out)
    print("wrote", out)
